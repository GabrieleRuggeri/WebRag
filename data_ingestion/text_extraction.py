'''
This module is intended to provide extractor classes for pdf, docx and txt documents.
'''

from typing import List, Dict
from abc import ABC, abstractmethod
import os
import logging
import pdfplumber
from docx import Document




class Extractor(ABC):
    """
    Abstract base class for document extractors.
    """

    @abstractmethod
    def extract_text(self, file_path: str) -> List[str]:
        """
        Extract text from a document.

        :param file_path: Path to the document.
        :return: List of extracted text strings.
        """
        pass

    @abstractmethod
    def extract_images(self, file_path: str) -> List[bytes]:
        """
        Extract images from a document.

        :param file_path: Path to the document.
        :return: List of extracted image bytes.
        """
        pass

class PdfExtractor(Extractor):
    """
    Extractor class for PDF documents.
    """

    def extract_text(self, file_path: str) -> List[str]:
        """
        Extract text from a PDF document.

        :param file_path: Path to the PDF document.
        :return: List of extracted text strings.
        """
        try:
            with pdfplumber.open(file_path) as pdf:
                text = []
                for page in pdf.pages:
                    text.append(page.extract_text())
                return text
        except Exception as e:
            logging.error(f"Error extracting text from PDF: {e}")
            return []
        
    def extract_images(self, file_path: str) -> List:
        """
        Extract images from a PDF document.

        :param file_path: Path to the PDF document.
        :return: List of extracted image bytes.
        """
        try:
            with pdfplumber.open(file_path) as pdf:
                images = []
                for page_num, page in enumerate(pdf.pages):
                    for img in page.images:
                        x0, y0, x1, y1 = img['bbox']
                        image = page.within_bbox((x0, y0, x1, y1)).to_image()
                        images.append((page_num, image.original))
                return images
        except Exception as e:
            logging.error(f"Error extracting images from PDF: {e}")
            return []
        
class DocxExtractor(Extractor):
    """
    Extractor class for DOCX documents.
    """

    def extract_text(self, file_path: str) -> List[str]:
        """
        Extract text from a DOCX document.

        :param file_path: Path to the DOCX document.
        :return: List of extracted text strings.
        """
        try:
            doc = Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                # add paragraoh titles
                if paragraph.style.name.startswith('Heading'):
                    text.append(f"Title: {paragraph.text}")
                else:
                    text.append(paragraph.text)
            return text
        except Exception as e:
            logging.error(f"Error extracting text from DOCX: {e}")
            return []
        
    def extract_images(self, file_path: str) -> List:
        """
        Extract images from a DOCX document.

        :param file_path: Path to the DOCX document.
        :return: List of extracted image bytes.
        """
        try:
            doc = Document(file_path)
            images = []
            for paragraph_num, rel in enumerate(doc.part.rels.values()):
                if "image" in rel.reltype:
                    image = rel.target_part.blob
                    images.append((paragraph_num, image))
            return images
        except Exception as e:
            logging.error(f"Error extracting images from DOCX: {e}")
            return []
    
class TxtExtractor(Extractor):
    """
    Extractor class for TXT documents.
    """

    def extract_text(self, file_path: str) -> List[str]:
        """
        Extract text from a TXT document.

        :param file_path: Path to the TXT document.
        :return: List of extracted text strings.
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.readlines()
            return text
        except Exception as e:
            logging.error(f"Error extracting text from TXT: {e}")
            return []
    def extract_images(self, file_path: str) -> List:
        return []

class DocumentExtractor:
    """
    Factory class to create document extractors based on file type.
    """

    def __init__(self):
        self.extractors = {
            'pdf': PdfExtractor(),
            'docx': DocxExtractor(),
            'txt': TxtExtractor()
        }
    def get_extractor(self, file_path: str) -> Extractor:
        """
        Get the appropriate extractor based on the file type.

        :param file_path: Path to the document.
        :return: Extractor instance.
        """
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()[1:]
        extractor = self.extractors.get(ext)
        if not extractor:
            raise ValueError(f"Unsupported file type: {ext}")
        return extractor
    
    def extract_text(self, file_path: str) -> List[str]:
        """
        Extract text from a document.
        :param file_path: Path to the document.
        :return: List of extracted text strings.
        """
        extractor = self.get_extractor(file_path)
        return extractor.extract_text(file_path)
    
    def extract_images(self, file_path: str) -> List:
        """
        Extract images from a document.
        :param file_path: Path to the document.
        :return: List of extracted image bytes.
        """
        extractor = self.get_extractor(file_path)
        return extractor.extract_images(file_path)
    
    def extract(self, file_path: str) -> Dict:
        """
        Extract text and images from a document.
        :param file_path: Path to the document.
        :return: Dictionary with extracted text and images.
        """
        extractor = self.get_extractor(file_path)
        text = extractor.extract_text(file_path)
        images = extractor.extract_images(file_path)
        return {
            'text': text,
            'images': images
        }
        
